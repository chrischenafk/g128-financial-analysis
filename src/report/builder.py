"""Run the skill's deterministic preprocessing locally, before the skill call.

The external skill's file mounting doesn't work over the API, so it otherwise
reconstructs the package from conversation text (losing the segmentation table,
charts, and full verification). To fix that, this builder runs the skill's own
``load_package.py`` and ``charts.py`` (vendored verbatim under ``engine/``) as
subprocesses and hands the resulting ``package.json`` + chart PNGs to the skill —
exactly the input it expects.

This is glue around the vendored scripts: it does no business computation and
never edits their logic. ``load_package.py`` failing is fatal (no package → no
report); ``charts.py`` failing is non-fatal (the report is still written, just
without charts). Deterministic, no network.
"""

from __future__ import annotations

import json
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path

from src import config
from src.utils.logger import get_logger

logger = get_logger(__name__)

_ENGINE = Path(__file__).resolve().parent / "engine"
LOAD_PACKAGE = _ENGINE / "load_package.py"
CHARTS = _ENGINE / "charts.py"

# Chart kinds requested from charts.py; each maps to "<kind>.png" in the outdir.
# charts.py emits only those whose data exists (e.g. trend needs history), so we
# collect from disk rather than assuming all three.
CHART_KINDS = ("bridge_mom", "bridge_yoy", "trend")


@dataclass
class ReportInputs:
    """Locally pre-processed inputs handed to the skill."""

    package_json: Path        # output of load_package.py
    charts: list[Path]        # PNGs that were actually generated (may be empty)
    workdir: Path             # the working directory these live in


def _read_period(package_dir: Path) -> str:
    """'YYYY-MM' period label from the package's run_metadata.json."""
    meta = json.loads((package_dir / "run_metadata.json").read_text(encoding="utf-8"))
    return str(meta["current_period"]["start"])[:7]  # "2026-04-01" → "2026-04"


def _run(cmd: list[str]) -> subprocess.CompletedProcess:
    """Run a subprocess, capturing stdout/stderr as text."""
    return subprocess.run(cmd, capture_output=True, text=True)


def _slim_package(package_json: Path, workdir: Path) -> Path:
    """Strip bulk raw arrays from package.json, keeping only what the skill
    needs for report writing. Returns path to the slimmed file.

    Removed (skill uses ranked{} subsets instead):
      - comparisons.mom / comparisons.yoy  — full 384-row arrays
      - sku_current                         — full 211-row array
      - known_dq_codes                      — static dict baked into the skill
      - supported_schema_versions           — housekeeping only

    Kept:
      - meta, channel (with bridges), anomalies, pipeline_warnings,
        loader_flags, context_md, ranked, historical, present, schema_version
    """
    pkg = json.loads(package_json.read_text(encoding="utf-8"))

    pkg.pop("sku_current", None)
    pkg.pop("known_dq_codes", None)
    pkg.pop("supported_schema_versions", None)
    # Keep ranked{} (pre-sorted top-N subsets) but drop the full comparison arrays
    # since ranked already contains the material movers the skill cites
    pkg.pop("comparisons", None)

    slim_path = workdir / "package_slim.json"
    slim_path.write_text(json.dumps(pkg), encoding="utf-8")  # no indent — saves ~20% vs indent=1

    size_before = package_json.stat().st_size
    size_after = slim_path.stat().st_size
    logger.info(
        "Slimmed package.json: %d KB → %d KB (removed raw SKU/comparison arrays).",
        size_before // 1024, size_after // 1024,
    )
    return slim_path


def prepare_report_inputs(package_dir: Path) -> ReportInputs:
    """Run load_package.py (fatal) then charts.py (best-effort) for ``package_dir``.

    Returns the ``package.json`` path plus the chart PNGs that were actually
    produced. Raises ``RuntimeError`` (naming the script + its stderr) if
    ``load_package.py`` fails; a ``charts.py`` failure is logged and yields an
    empty charts list rather than aborting.
    """
    period = _read_period(package_dir)
    workdir = config.OUTPUT_REPORTS / ".build" / period
    charts_dir = workdir / "charts"
    workdir.mkdir(parents=True, exist_ok=True)
    package_json = workdir / "package.json"

    # Step 1 — load_package.py (fatal on failure: no package → no report).
    result = _run([sys.executable, str(LOAD_PACKAGE), str(package_dir), "-o", str(package_json)])
    if result.returncode != 0:
        raise RuntimeError(
            f"load_package.py failed (exit {result.returncode}): "
            f"{result.stderr.strip() or '<no stderr>'}"
        )
    if result.stdout.strip():
        logger.info("load_package.py: %s", result.stdout.strip())

    # Step 1b — slim package.json for upload (drop bulky raw arrays). Charts still
    # read the full package.json below; only the uploaded copy is slimmed.
    slim_json = _slim_package(package_json, workdir)

    # Step 2 — charts.py (best-effort: a failure must not block the report).
    charts: list[Path] = []
    chart_result = _run([
        sys.executable, str(CHARTS), str(package_json),
        "--outdir", str(charts_dir), "--which", ",".join(CHART_KINDS),
    ])
    if chart_result.returncode != 0:
        logger.warning(
            "charts.py failed (exit %d): %s — continuing without charts.",
            chart_result.returncode, chart_result.stderr.strip() or "<no stderr>",
        )
    else:
        # Collect only the PNGs that actually exist (e.g. trend is skipped with no history).
        charts = [charts_dir / f"{kind}.png" for kind in CHART_KINDS
                  if (charts_dir / f"{kind}.png").exists()]
        logger.info("charts.py produced %d chart(s): %s", len(charts), [p.name for p in charts])

    # Hand the skill the slimmed package.json (charts were rendered from the full one).
    return ReportInputs(package_json=slim_json, charts=charts, workdir=workdir)


# ─────────────────────────────────────────────────────────────────────────────
# Post-skill: inject the real charts into the downloaded .docx
# ─────────────────────────────────────────────────────────────────────────────
def _inject_charts(docx_path: Path, charts: list[Path]) -> Path:
    """Insert the real chart PNGs before their anchor paragraphs in the .docx.

    build_doc.js skips the chart images (the files aren't in its container), so the
    document has no embedded images to swap — instead each section carries a text
    anchor naming its chart (e.g. a paragraph containing "bridge_mom.png"). This
    inserts an image paragraph immediately before each matching anchor, using the
    locally-rendered PNG.

    Best-effort: returns ``docx_path`` regardless; if no anchors are found it logs a
    WARNING and skips. python-docx is imported lazily, so a missing install (or any
    failure, via main.py's try/except) degrades to "no charts" rather than breaking
    the run.
    """
    if not charts:
        return docx_path

    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches

    # filename stem → anchor text variants to search for. The skill references the
    # charts inconsistently (with/without ".png", or in prose like "MoM bridge"),
    # so we match any variant rather than the literal filename alone.
    # The generic "Profit bridge:" / "profit bridge" variants live ONLY on
    # bridge_mom: the skill captions the bridge as "Profit bridge: <from> → <to>"
    # (no MoM/YoY token). In a MoM-only run that bare caption is the sole bridge,
    # so bridge_mom must match it. In a two-file run the MoM section precedes the
    # YoY one, so the FIRST "Profit bridge:" is the MoM caption — bridge_mom claims
    # it correctly. bridge_yoy deliberately keeps only its YoY-specific variants so
    # the two stems can't both anchor to that same first caption.
    ANCHORS = {
        "bridge_mom": [
            "bridge_mom.png", "bridge_mom",
            "MoM Profit Bridge", "MoM bridge",
            "Profit bridge:", "profit bridge",
        ],
        "bridge_yoy": ["bridge_yoy.png", "bridge_yoy", "YoY Profit Bridge", "YoY bridge"],
    }

    doc = Document(str(docx_path))

    # Strip the skill's placeholder images before injecting the real charts.
    # build_doc.js embeds placeholder PNGs in the body (it can't mount the chart
    # files); left in place, our injected charts would sit *beside* them, doubling
    # every figure. Remove only <w:drawing> elements in the document body — headers,
    # footers, and text boxes live outside doc.element.body, so a logo or page
    # furniture is untouched. Removing just the drawing leaves the paragraph (and
    # any caption text) intact, so the anchor search below still matches correctly.
    body = doc.element.body
    drawings = body.findall(".//" + qn("w:drawing"))
    for drawing in drawings:
        drawing.getparent().remove(drawing)
    logger.info("Stripped %d placeholder image(s) from body before chart injection.",
                len(drawings))

    # Build a lookup: chart stem -> paragraph index of its anchor (first match wins).
    anchor_map: dict[str, int] = {}
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        for stem, variants in ANCHORS.items():
            if stem not in anchor_map and any(v in text for v in variants):
                anchor_map[stem] = i

    if not anchor_map:
        logger.warning("_inject_charts: no anchor paragraphs found in %s — skipping.", docx_path.name)
        return docx_path

    # Positional fallback: the skill occasionally references a chart outside its
    # section — too early (a stray mention in the cover-page caveats, before the
    # Executive Summary) or too late (an appendix). A valid section anchor sits
    # between paragraph 15 (the earliest the Executive Summary could start) and
    # paragraph 80 (after which we're into the appendix). An anchor below 15 or
    # above 80 is almost certainly misplaced — re-anchor it to just after the
    # relevant section heading instead. The heading strings are stable because
    # section numbering is locked in the trigger message. If the anchor is already
    # in its section, this never fires.
    HEADING_FALLBACKS = {
        "bridge_mom": "3. MoM Performance",
        "bridge_yoy": "4. YoY / Seasonality Context",
    }
    for stem, heading_text in HEADING_FALLBACKS.items():
        if stem in anchor_map and (anchor_map[stem] < 15 or anchor_map[stem] > 80):
            misplacement = "cover page" if anchor_map[stem] < 15 else "appendix"
            for i, para in enumerate(doc.paragraphs):
                if heading_text in para.text:
                    logger.warning(
                        "_inject_charts: %s anchor at paragraph %d is outside the valid "
                        "section range (likely %s) — falling back to heading position %d ('%s').",
                        stem, anchor_map[stem], misplacement, i, heading_text,
                    )
                    anchor_map[stem] = i + 1  # inject right after the heading
                    break
            else:
                logger.warning(
                    "_inject_charts: %s anchor at paragraph %d is outside the valid "
                    "section range (likely %s) but heading '%s' was not found — leaving "
                    "anchor in place. Section wording may have drifted.",
                    stem, anchor_map[stem], misplacement, heading_text,
                )

    # Diagnose a partial match: chart rendered locally but no anchor for it in the doc.
    for chart in charts:
        if chart.stem in ANCHORS and chart.stem not in anchor_map:
            logger.warning(
                "_inject_charts: chart %s exists but no anchor found for it in %s — "
                "it will not be injected.", chart.name, docx_path.name,
            )

    chart_map = {p.stem: p for p in charts}
    # One insertion per stem (anchor_map already enforces first-match-wins above).
    # Insert in ASCENDING index order, tracking an offset for the index shift each
    # insertion introduces — every inserted paragraph pushes all later anchors down
    # by one, so the next anchor's live index is its original index + offset.
    insertions = sorted(
        [(anchor_map[stem], chart_map[stem]) for stem in anchor_map if stem in chart_map],
        key=lambda x: x[0],
    )

    offset = 0  # number of paragraphs inserted so far → shift for subsequent anchors
    for para_idx, chart_path in insertions:
        adjusted_idx = para_idx + offset
        anchor_para = doc.paragraphs[adjusted_idx]
        # Insert a fresh, center-aligned paragraph before the anchor to hold the image.
        new_para = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        pPr.append(jc)
        new_para.append(pPr)
        anchor_para._element.addprevious(new_para)

        img_para = doc.paragraphs[adjusted_idx]  # the new empty paragraph
        run = img_para.add_run()
        run.add_picture(str(chart_path), width=Inches(6.0))
        logger.info("Injected %s before paragraph %d.", chart_path.name, para_idx)
        offset += 1  # each insertion shifts all subsequent anchors by one

    doc.save(str(docx_path))
    logger.info("Chart injection complete: %s", docx_path.name)
    return docx_path
